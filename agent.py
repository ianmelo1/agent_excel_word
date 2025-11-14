import os
import json
import platform
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rate_limiter import rate_limit
import google.generativeai as genai


class AgenteOfficeIA:
    """
    Agente automático que integra Excel, Word e IA (Gemini)
    """

    def __init__(self, api_key=None, modelo="gemini-2.0-flash"):
        """
        Inicializa o agente com a chave da API do Google Gemini

        Args:
            api_key: Chave da API do Google
            modelo: Nome do modelo Gemini (padrão: gemini-2.0-flash-exp)
        """
        self.api_key = api_key or os.environ.get("GOOGLE_API_KEY")
        self.modelo = modelo

        if self.api_key:
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel(modelo)
            print(f"✅ Gemini inicializado: {modelo}")
        else:
            self.model = None
            print("⚠️  API Key não fornecida. Funções de IA estarão desabilitadas.")

    # ============ FUNÇÕES EXCEL ============

    def criar_excel(self, arquivo, dados, cabecalhos=None):
        """
        Cria um arquivo Excel com dados e formatação

        Args:
            arquivo: nome do arquivo .xlsx
            dados: lista de listas com os dados
            cabecalhos: lista com nomes das colunas
        """
        wb = openpyxl.Workbook()
        ws = wb.active

        # Adiciona cabeçalhos se fornecidos
        if cabecalhos:
            ws.append(cabecalhos)
            # Formata cabeçalhos
            for cell in ws[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                cell.alignment = Alignment(horizontal="center")

        # Adiciona dados
        for linha in dados:
            ws.append(linha)

        # Ajusta largura das colunas
        for column in ws.columns:
            max_length = 0
            column = [cell for cell in column]
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column[0].column_letter].width = adjusted_width

        wb.save(arquivo)
        print(f"✅ Excel criado: {arquivo}")
        return arquivo

    def ler_excel(self, arquivo, sheet=None):
        """
        Lê dados de um arquivo Excel

        Args:
            arquivo: nome do arquivo .xlsx
            sheet: nome da planilha (opcional)

        Returns:
            Lista de listas com os dados
        """
        wb = openpyxl.load_workbook(arquivo)
        ws = wb[sheet] if sheet else wb.active

        dados = []
        for row in ws.iter_rows(values_only=True):
            dados.append(list(row))

        print(f"✅ Excel lido: {arquivo} ({len(dados)} linhas)")
        return dados

    def atualizar_excel(self, arquivo, linha, coluna, valor):
        """
        Atualiza uma célula específica do Excel
        """
        wb = openpyxl.load_workbook(arquivo)
        ws = wb.active
        ws.cell(row=linha, column=coluna, value=valor)
        wb.save(arquivo)
        print(f"✅ Excel atualizado: célula ({linha},{coluna}) = {valor}")

    # ============ FUNÇÕES WORD ============

    def criar_word(self, arquivo, titulo, conteudo):
        """
        Cria um documento Word formatado

        Args:
            arquivo: nome do arquivo .docx
            titulo: título do documento
            conteudo: texto ou lista de parágrafos
        """
        doc = Document()

        # Adiciona título
        heading = doc.add_heading(titulo, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adiciona data
        data_para = doc.add_paragraph()
        data_run = data_para.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        data_run.font.size = Pt(9)
        data_run.font.color.rgb = RGBColor(128, 128, 128)
        data_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()  # Espaço

        # Adiciona conteúdo
        if isinstance(conteudo, list):
            for paragrafo in conteudo:
                doc.add_paragraph(paragrafo)
        else:
            doc.add_paragraph(conteudo)

        doc.save(arquivo)
        print(f"✅ Word criado: {arquivo}")
        return arquivo

    def ler_word(self, arquivo):
        """
        Lê o conteúdo de um documento Word
        """
        doc = Document(arquivo)
        conteudo = []

        for para in doc.paragraphs:
            if para.text.strip():
                conteudo.append(para.text)

        print(f"✅ Word lido: {arquivo} ({len(conteudo)} parágrafos)")
        return conteudo

    def adicionar_ao_word(self, arquivo, texto):
        """
        Adiciona conteúdo a um documento Word existente
        """
        doc = Document(arquivo)
        doc.add_paragraph(texto)
        doc.save(arquivo)
        print(f"✅ Conteúdo adicionado ao Word: {arquivo}")

    # ============ FUNÇÕES IA ============

    @rate_limit(max_per_minute=10)
    def perguntar_ia(self, pergunta, contexto=None):
        """
        Faz uma pergunta para a IA Gemini

        Args:
            pergunta: pergunta ou comando
            contexto: informação adicional para contexto

        Returns:
            Resposta da IA
        """
        if not self.model:
            return "Erro: API Key não configurada"

        prompt = pergunta
        if contexto:
            prompt = f"Contexto: {contexto}\n\nPergunta: {pergunta}"

        try:
            response = self.model.generate_content(prompt)
            resposta = response.text
            print(f"✅ IA respondeu ({len(resposta)} caracteres)")
            return resposta
        except Exception as e:
            return f"❌ Erro ao consultar IA: {str(e)}"

    def analisar_excel_com_ia(self, arquivo):
        """
        Lê um Excel e pede para IA analisar os dados
        """
        dados = self.ler_excel(arquivo)

        prompt = f"""Analise os seguintes dados de uma planilha Excel:

{json.dumps(dados[:10], ensure_ascii=False)}

Forneça:
1. Um resumo dos dados
2. Insights principais
3. Sugestões de análise"""

        return self.perguntar_ia(prompt)

    # ============ FUNÇÕES AUTOMÁTICAS ============

    def relatorio_automatico(self, dados_excel, arquivo_saida="relatorio.docx"):
        """
        Cria um relatório Word automático baseado em dados do Excel
        """
        # Analisa os dados com IA
        analise = self.perguntar_ia(
            f"Crie um relatório executivo baseado nestes dados: {json.dumps(dados_excel[:5], ensure_ascii=False)}"
        )

        # Cria o Word
        self.criar_word(
            arquivo_saida,
            "Relatório Automatizado",
            analise
        )

        return arquivo_saida

    def pipeline_completo(self, dados, nome_projeto="projeto"):
        """
        Executa um pipeline completo: Excel -> IA -> Word
        """
        print(f"\n🚀 Iniciando pipeline: {nome_projeto}")

        # 1. Cria Excel
        arquivo_excel = f"{nome_projeto}.xlsx"
        self.criar_excel(
            arquivo_excel,
            dados,
            cabecalhos=["ID", "Descrição", "Valor", "Status"]
        )

        # 2. Analisa com IA
        print("\n🤖 Analisando dados com IA...")
        analise = self.analisar_excel_com_ia(arquivo_excel)

        # 3. Cria relatório Word
        arquivo_word = f"{nome_projeto}_relatorio.docx"
        self.criar_word(
            arquivo_word,
            f"Relatório: {nome_projeto}",
            [
                "Este relatório foi gerado automaticamente pelo agente.",
                "",
                "ANÁLISE DOS DADOS:",
                analise
            ]
        )

        print(f"\n✨ Pipeline concluído!")
        print(f"   📊 Excel: {arquivo_excel}")
        print(f"   📄 Word: {arquivo_word}")

        return arquivo_excel, arquivo_word


# ============ FUNÇÃO AUXILIAR ============

def abrir_arquivo(arquivo):
    """Abre arquivo no programa padrão do sistema operacional"""
    if platform.system() == 'Windows':
        os.startfile(arquivo)
    elif platform.system() == 'Darwin':  # Mac
        os.system(f'open "{arquivo}"')
    else:  # Linux
        os.system(f'xdg-open "{arquivo}"')
    print(f"📂 Abrindo arquivo: {arquivo}")


# ============ EXEMPLOS DE USO ============

if __name__ == "__main__":
    # Inicializa o agente
    agente = AgenteOfficeIA()  # Ou: AgenteOfficeIA(api_key="sua-chave", modelo="gemini-1.5-pro")

    print("=" * 60)
    print("🤖 AGENTE PYTHON - EXCEL, WORD E IA (GEMINI)")
    print("=" * 60)

    # Exemplo 1: Criar Excel simples
    print("\n📊 Exemplo 1: Criando Excel...")
    dados_vendas = [
        [1, "Produto A", 1500, "Concluído"],
        [2, "Produto B", 2300, "Pendente"],
        [3, "Produto C", 1800, "Concluído"],
        [4, "Produto D", 3200, "Em Análise"],
        [5, "Produto E", 900, "Concluído"]
    ]

    agente.criar_excel(
        "vendas.xlsx",
        dados_vendas,
        cabecalhos=["ID", "Produto", "Valor (R$)", "Status"]
    )
    abrir_arquivo("vendas.xlsx")  # Abre automaticamente

    # Exemplo 2: Criar Word
    print("\n📄 Exemplo 2: Criando Word...")
    agente.criar_word(
        "relatorio.docx",
        "Relatório de Vendas",
        [
            "Este é um relatório automático gerado pelo agente Python.",
            "",
            "Os dados foram processados e formatados automaticamente.",
            "Integração completa com Excel e IA."
        ]
    )
    abrir_arquivo("relatorio.docx")  # Abre automaticamente

    # Exemplo 3: Pipeline completo (requer API key)
    print("\n🚀 Exemplo 3: Pipeline completo...")
    print("   (Configure ANTHROPIC_API_KEY para usar IA)")

    # Descomente para executar com IA:
    # agente.pipeline_completo(dados_vendas, "vendas_q4")

    print("\n✅ Exemplos concluídos!")
    print("\nPara usar IA, configure:")
    print("export GOOGLE_API_KEY='sua-chave-aqui'")
    print("\nOu passe no construtor:")
    print("agente = AgenteOfficeIA(api_key='sua-chave')")
    print("\n📋 Modelos disponíveis:")
    print("   - gemini-2.0-flash-exp (recomendado - mais novo)")
    print("   - gemini-1.5-flash (rápido)")
    print("   - gemini-1.5-pro (mais inteligente)")
    print("\nExemplo: AgenteOfficeIA(api_key='...', modelo='gemini-1.5-pro')")